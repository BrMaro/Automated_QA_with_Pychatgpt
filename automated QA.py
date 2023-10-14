from pyChatGPT import ChatGPT
from pdfminer.high_level import extract_text
from tkinter import *
from tkinter import filedialog,messagebox
import re
import os
import math
import tempfile
import docx
import time
from docx import Document

session_token = input("Copy and paste your session token from ChatGPT website")
api = ChatGPT(session_token)


def split_pdf(file_path, chunk_size):
    total_pages = len(list(extract_text(file_path)))
    if total_pages <= chunk_size:
        return [file_path]
    num_chunks = math.ceil(total_pages / chunk_size)
    temp_dir = tempfile.mkdtemp()
    file_paths = []
    for i in range(num_chunks):
        start = i * chunk_size
        end = min((i + 1) * chunk_size, total_pages)
        pages = list(range(start + 1, end + 1))
        temp_path = os.path.join(temp_dir, f"chunk_{i + 1}.pdf")
        os.system(f"pdfseparate -f {start + 1} -l {end} {file_path} {temp_path}")
        file_paths.append(temp_path)
    return file_paths


window = Tk()
window.title("Automated Question and Answer Generator")
window.geometry("1000x600")
window.configure(bg = "#000000")


def generate_questions_and_answers(num_questions, topic):

    resp = api.send_message([f"Generate {num_questions} questions based on the {topic} topic in the above text"])
    questions = resp['message']
    print(questions)
    resp = api.send_message(["Generate the answers to the above questions in their respective order"])
    answers = resp['message']
    print(answers)
    document = Document()
    document.add_heading('Questions', 0)
    document.add_paragraph(questions)
    document.add_heading('Answers',0)
    document.add_paragraph(answers)
    document.save('document.docx')
    os.startfile('document.docx')


def btn_clicked():
    WORD_COUNT_THRESHOLD = 1000
    b0.destroy()
    file_path = filedialog.askopenfilename()
    file = ''
    while os.path.splitext(file_path)[1] != '.pdf' and os.path.splitext(file_path)[1] != '.docx':
        messagebox.showerror("Invalid file type", "The selected file is neither a pdf nor a word file.")
        file_path = filedialog.askopenfilename()
    if os.path.splitext(file_path)[1] == '.pdf':
        file = extract_text(file_path)
    if os.path.splitext(file_path)[1] == '.docx':
        doc = docx.Document(file_path)
        file = ''
        for para in doc.paragraphs:
            file += para.text
    doc = re.sub('\s+', ' ', file)
    word_count = len(doc.split())
    messagebox.showinfo("Word_count", f"The file has {word_count} words")

    if word_count > WORD_COUNT_THRESHOLD:
        file_paths=split_pdf(file_path,WORD_COUNT_THRESHOLD)
        titles = []
        for chunk in file_paths:
            chunk = extract_text(chunk)
            resp = api.send_message([f"What are the headings in the text below in bullet form: {chunk}"])
            topics = resp['message']
            titles.append(topics)
        messagebox.showinfo("Topics or Headings",f"The selected file has the following headings or topics {titles}")
    else:
            resp = api.send_message([f"What are the headings in the text below in bullet form: {doc}"])
            titles = resp['message']
            messagebox.showinfo("Topics or Headings",
                                f"The selected file has the following headings or topics {titles}")

    time.sleep(2)
    num = canvas.create_text(
        499.5, 270.0,
        text="Enter number of questions you want to generate: ",
        fill="#ffffff",
        font=("Inter-Light", int(13.0)))

    e = Entry(window, width=3)
    e.place(x=680, y=260)

    topic_text = canvas.create_text(
        499.5, 300.0,
        text=f"Enter the desired topic from the options provided previously and press 'Enter': ",
        fill="#ffffff",
        font=("Inter-Light", int(12.0)))
    topic_input = Entry(window, width=20)
    topic_input.place(x=770, y=290)

    topic_input.bind('<Return>', lambda event: generate_questions_and_answers(e.get(), topic_input.get()))

    time.sleep(15)






canvas = Canvas(
    window,
    bg = "#000000",
    height = 800,
    width = 1200,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge")
canvas.place(x = 0, y = 0)

canvas.create_text(
    499.5, 131.5,
    text = "AUTOMATIC QUESTION & ANSWER GENERATOR",
    fill = "#ffffff",
    font = ("Inter-Light", int(30.0)))

canvas.create_text(
    499.5, 220.0,
    text = "Smart solutions made easy and accessible",
    fill = "#ffffff",
    font = ("Inter-Light", int(16.0)))


img0 = PhotoImage(file = f"img0.png")
b0 = Button(
    image = img0,
    borderwidth = 0,
    highlightthickness = 0,
    command = btn_clicked,
    relief = "flat")

b0.place(
    x = 423, y = 377,
    width = 154,
    height = 45)

window.resizable(False, False)
window.mainloop()









